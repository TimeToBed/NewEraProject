# -*- coding: utf-8 -*-
from docx import Document
import re, os
import json
from tqdm import tqdm
from .knowledge import LLm_KnowledgeBase

split_regular = {
    '小题号':r'\d+[．\.]\s*',
    '中题号':r'（[一二三四五六七八九十]）',
    '大题号':r'[一二三四五六七八九十]、\s*',
    '子题号':r'（\d）',
    'ocr小题号':r'\b([0-9]|1[0-9]|2[0-2])\b'
}

def trans_key_values(exam_js_path, key):
    def deal(x:str):
        out = x.replace("```json", "").replace("```", "").replace("\n","").strip()
        output = json.loads(out)
        return output
    
    with open(exam_js_path, 'r', encoding='utf-8') as exam_json:
        exam_info = json.load(exam_json)
    
    for info_k, info_v in exam_info.items():
        if info_v['sub_topic']:
            for sub_k, sub_v in info_v['sub_topic'].items():
                if key in sub_v:
                    sub_v[key] = deal(sub_v[key])
        else:
            if key in info_v:
                info_v[key] = deal(info_v[key])

    with open(exam_js_path, 'w', encoding='utf-8') as json_file:
        json.dump(exam_info, json_file, ensure_ascii=False)


def sanitize_filename(text):
    return re.sub(r'[\/:*?"<>|（），]', '', text)


#将一段字符串列表按照pattern进行划分
def split_by_pattern(lines_list, pattern): 
    """将一段字符串列表按照pattern进行划分

    args:
        -lines_list: 字符串列表，每一个元素代表一行内容
        -pattern: 模式匹配字符串
    output:
        -content_dict: 字典，键为要匹配的字符串，值为两个匹配字符串间的内容
    """
    current_key = None
    content_dict,content_list = {}, []
    for line in lines_list:
        match = re.match(pattern, line)
        if match:
            current_key =  match.group(1)
            content_dict[current_key] = [line]
        elif current_key is not None:
            content_dict[current_key].append(line)
        else:
            pass #即无可匹配内容，也无存在的键。

    # 将每个题目的段落列表合并成一个字符串
    for key in content_dict:
        content_dict[key] = '\n'.join(content_dict[key])
    

    for line in lines_list:
        match = re.match(pattern, line)
        if match:
            content_list.append([line])
        elif len(content_list)>0:
            content_list[-1].append(line)
        else:
            pass  #即无可匹配内容，也无存在的列表。
    
    return content_dict, content_list


#获得题目字典，给予每道题目上下文信息，用于遍历初版exam的每道题目
def get_topic_dict(content_dict):
    """
    获得题目字典，给予每道题目上下文信息。用作大语言模型分析
    output:
        {
            "题号": {              #题号为数字 1 2 3...
                "content": "....", #该题目内容
                "pretext": "....", #该题目的上下文信息，主要为材料
                "score":"...",     #该题目的分值
                "sub_topic": null  #是否有子题目，比如（1）（2）等，该值为同形式字典，但该字典中无"sub_topic"键
                },
        }
    """
    topic_dict = {}
    pretext = None
    #构建小题号键和值,此时大题号和中题号都被和为小题号的上下文
    for key in content_dict:

        if key[:-1].isdigit():
            topic_dict[key[:-1]] = {"content":content_dict[key]}
            if pretext is not None:
                topic_dict[key[:-1]].update({"pretext":pretext})
        else:
            pretext = content_dict[key]
    
    #构建小题号包含的子题号键值
    for topic_k,topic_v in topic_dict.items():
        sub_dict = {}
        pretext = None
        content_lines = topic_v["content"].split('\n')
        pattern = f"^({split_regular['子题号']}|{split_regular['小题号']})"
        dicts,_ = split_by_pattern(content_lines, pattern)

        if len(dicts) > 1: #存在子题号

            for key in dicts:
                if not key[:-1].isdigit():
                    sub_dict[key] = {"content":dicts[key]}
                    if pretext is not None:
                        sub_dict[key].update({"pretext":pretext})
                    score = re.search(r"（\d+分）",sub_dict[key]["content"])
                    if score is not None:
                        sub_dict[key].update({'score':score.group()[1:-2]})
                else:
                    pretext = dicts[key]

            topic_dict[topic_k].update({'sub_topic':sub_dict})
        else:
            score = re.search(r"（\d+分）",topic_dict[topic_k]["content"])
            if score is not None:
                topic_dict[topic_k].update({'score':score.group()[1:-2]})
            topic_dict[topic_k].update({'sub_topic':None})
        
    return topic_dict

#获得ocr中每个chunk(根据中题号划分)中题目的内容和答案
def get_ocr_content_dict(chunk_dict):
    """
    获得题目字典，给予每道题目上下文信息。用作大语言模型分析
    output:
        {
            "题号": {              #题号为数字 1 2 3...
                "content": "....", #该题目内容},
        }
    """
    ocr_dict = {}
    #构建小题号键和值,此时大题号和中题号都被和为小题号的上下文
    for key in chunk_dict:

        if key.isdigit():
            ocr_dict[key] = {"content":chunk_dict[key]}
    
    #构建小题号包含的子题号键值
    for topic_k,topic_v in ocr_dict.items():
        sub_dict = {}
        content_lines = topic_v["content"].split('\n')
        pattern = f"^({split_regular['子题号']}|{split_regular['ocr小题号']})"
        dicts,_ = split_by_pattern(content_lines, pattern)

        if len(dicts) > 1: #存在子题号

            for key in dicts:
                if not key.isdigit():
                    sub_dict[key] = {"content":dicts[key]}

            ocr_dict[topic_k] = sub_dict
        
    return ocr_dict

#获得ocr中每个chunk(根据中题号划分)中题目的内容和答案
def get_framework_dict(chunk_dict):
    """
    获得题目字典，给予每道题目上下文信息。用作大语言模型分析
    output:
        {
            "题号": {              #题号为数字 1 2 3...
                "content": "....", #该题目内容},
        }
    """
    ocr_dict = {}
    #构建小题号键和值,此时大题号和中题号都被和为小题号的上下文
    for key in chunk_dict:

        if key.isdigit():
            ocr_dict[key] = {"content":chunk_dict[key]}
    
    #构建小题号包含的子题号键值
    for topic_k,topic_v in ocr_dict.items():
        sub_dict = {}
        content_lines = topic_v["content"].split('\n')
        pattern = f"^({split_regular['子题号']}|{split_regular['ocr小题号']})"
        dicts,_ = split_by_pattern(content_lines, pattern)

        if len(dicts) > 1: #存在子题号

            for key in dicts:
                if not key.isdigit():
                    sub_dict[key] = {"content":dicts[key]}

            ocr_dict[topic_k] = sub_dict
        
    return ocr_dict

#将ocr文件正常分割
def load_ocr(json_path):

    with open(json_path, 'r', encoding='utf-8') as json_file:
    # 使用 json.load() 方法，将读取的文件内容转换成一个字典
        data = json.load(json_file)

    select = data.pop('select').split('\n')
    record_match = re.search(r'学号：([^姓名]*)姓名：([^\\n]*)', data['1'].split('\n')[0])

    _,answer_list = split_by_pattern(select, split_regular['ocr小题号'])

    page_dict, content_dict, all_content = {}, {}, ''
    for k,v in data.items():
        all_content += v
        lines = v.split('\n')
        for line in lines:
            match = re.match(split_regular['ocr小题号'], line)

            if match:
                page_dict[match.group()] = {'page':int(k)}


    pattern = f"^({split_regular['中题号']})"
    lines_list = all_content.split('\n')
    _,chunk_list = split_by_pattern(lines_list, pattern)

    pattern = f"^({split_regular['ocr小题号']}|{split_regular['中题号']})"
    for v in chunk_list:

        item_dict,_ = split_by_pattern(v,pattern)
        dict_data = get_ocr_content_dict(item_dict)
        content_dict.update(dict_data)
    
    for k,v in content_dict.items():
        if len(content_dict[k].keys()) > 1:
            for k_sub in content_dict[k]:
                content_dict[k][k_sub].update(page_dict[k])
        else:
            content_dict[k].update(page_dict[k])
    
    for it in answer_list:
        content_dict[it[0]]['content'] += '学生答案'+it[1]

    #with open(json_path, 'w', encoding='utf-8') as json_file:
    #    json.dump(paper_dict, json_file, ensure_ascii=False)
    content_dict.update({'学号':re.sub(r'\D', '', record_match.group(1)),'姓名':record_match.group(2)})

    return content_dict

#
def init_llm_json(exam_doc, ocr_path, analysis_path, save_path):

    def update_dict(dict_obj, k_s, ik_sub = None):
        if ik_sub:
            dict_obj.update({
                'analysis': analysis_json[k_s]['sub_topic'][ik_sub]['analysis']['analysis'],
                'score': analysis_json[k_s]['sub_topic'][ik_sub]['score'],
                'content': analysis_json[k_s]['sub_topic'][ik_sub]['content'],
                'ocr_content': ocr_json[k_s][ik_sub]['content'],
                'page': ocr_json[k_s][ik_sub]['page'],
                'llm_mark': None,
                'llm_comment': '',
                'mark_score': None,
                'mark': ''
            })
        else:
            dict_obj.update({
                'analysis': analysis_json[k_s]['analysis']['analysis'],
                'score': analysis_json[k_s]['score'],
                'content': analysis_json[k_s]['content'],
                'ocr_content': ocr_json[k_s]['content'],
                'page': ocr_json[k_s]['page'],
                'llm_mark': None,
                'llm_comment': '',
                'mark_score': None,
                'mark': ''
            })
            
        return dict_obj
    
    #读取ocr的处理后json文件
    ocr_json = load_ocr(ocr_path)
    #读取LLM与分析json文件
    with open(analysis_path, 'r', encoding='utf-8') as json_file:
        analysis_json = json.load(json_file)
    #读取考试试卷doc文档
    document = Document(exam_doc)

    frame_dict = {'学号':ocr_json['学号'],'姓名':ocr_json['姓名']}

    lines_list = [paragraph.text for paragraph in document.paragraphs]
    chunks_dict, _ = split_by_pattern(lines_list, f"^({split_regular['大题号']})")
    for k_c, v_c in chunks_dict.items():
        key = k_c[:-1]
        frame_dict[key] = {}
        medium_dict, _ = split_by_pattern(v_c.split('\n'), f"^({split_regular['中题号']})")

        for k_m, v_m in medium_dict.items() if medium_dict else [(None, v_c)]:
            if k_m:
                frame_dict[key][k_m] = {}
            small_dict, _ = split_by_pattern(v_m.split('\n') if k_m else v_c.split('\n'),
                                             f"^({split_regular['小题号']})")
            for k_sr, v_s in small_dict.items():
                k_s = k_sr[:-1]
                if k_m:
                    frame_dict[key][k_m][k_s] = {}
                    target_dict = frame_dict[key][k_m][k_s]
                else:
                    frame_dict[key][k_s] = {}
                    target_dict = frame_dict[key][k_s]
                if analysis_json[k_s]['sub_topic']:
                    for ik_sub in analysis_json[k_s]['sub_topic']:
                        target_dict[ik_sub] = {}
                        target_dict[ik_sub] = update_dict(target_dict[ik_sub], k_s, ik_sub)
                else:
                    target_dict = update_dict(target_dict, k_s)

    with open(save_path, 'w', encoding='utf-8') as f:
        json.dump(frame_dict, f, ensure_ascii=False)
                    
    return frame_dict
            
#获得初版exam
def load_exam(doc_path, json_path=None):

    document = Document(doc_path)  # 打开文件example.docx

    pattern = f"^({split_regular['大题号']})"
    lines_list = [paragraph.text for paragraph in document.paragraphs]
    chunks_dict,_ = split_by_pattern(lines_list, pattern)

    paper_dict = {}
    pattern = f"^({split_regular['小题号']}|{split_regular['中题号']}|{split_regular['大题号']})"
    for k,v in chunks_dict.items():
        item_dict,_ = split_by_pattern(v.split('\n'),pattern)
        dict_data = get_topic_dict(item_dict)
        paper_dict.update(dict_data)

    #with open(json_path, 'w', encoding='utf-8') as json_file:
    #    json.dump(paper_dict, json_file, ensure_ascii=False)
    return paper_dict

def load_answer(answer_path, save_json_path=None):

    document = Document(answer_path)

    pattern = f"^({split_regular['小题号']}|{split_regular['子题号']})"
    lines_list = [paragraph.text for paragraph in document.paragraphs]
    items_dict, item_list = split_by_pattern(lines_list, pattern)
    
    answer_dict = {}

    current_k = None
    for item in item_list:
        match = re.match(split_regular['小题号'],item[0])
        match_sub = re.match(split_regular['子题号'],item[0])
        if match:
            key = match.group()[:-1]
            answer_dict[key] = item[0][len(key)+1:]
            current_k = key
        elif match_sub:
            key = match_sub.group()
            if isinstance(answer_dict[current_k], dict):
                answer_dict[current_k].update({key:item[0][len(key):]})
            else:
                answer_dict[current_k] = {key:item[0][len(key):]}
    
    #with open(save_json_path, 'w', encoding='utf-8') as json_file:
    #    json.dump(answer_dict, json_file, ensure_ascii=False)
    return answer_dict

def add_answer_to_exam(exam_js, answer_js):
    """
    获得题目字典，给予每道题目上下文信息。用作大语言模型分析
    output:
        {
            "题号": {              #题号为数字 1 2 3...
                "content": "....", #该题目内容
                "pretext": "....", #该题目的上下文信息，主要为材料
                "sub_topic": null  #是否有子题目，比如（1）（2）等，该值为同形式字典，但该字典中无"sub_topic"键无"answer"键
                "answer": "..."
                },
        }
    """
    #with open(exam_js_path, 'r', encoding='utf-8') as paper_json:
    #    paper = json.load(paper_json)
    #with open(answer_js_path, 'r', encoding='utf-8') as answer_json:
    #    answer = json.load(answer_json)
    paper, answer = exam_js, answer_js
    for paper_k, paper_v in paper.items():
        
        if paper_k in answer:

            if paper_v["sub_topic"]:
                for sub_k, sub_v in paper_v["sub_topic"].items():
                    if sub_k in answer[paper_k]:
                        paper[paper_k]["sub_topic"][sub_k].update({'answer':answer[paper_k][sub_k]})
                    else:
                        paper[paper_k]["sub_topic"][sub_k].update({'answer':"略"})
                paper[paper_k].update({'answer':"略"})
            else:
                paper[paper_k].update({'answer':answer[paper_k]})
        else:
            paper[paper_k].update({'answer':"略"})
    
    #with open('paper.json', 'w', encoding='utf-8') as json_file:
    #    json.dump(paper, json_file, ensure_ascii=False)
    return paper

def simplify_content(exam_info,save_js_path,kdb_path):
    
    print(f" 知识库路径 {kdb_path}")
    
    llm_kb = LLm_KnowledgeBase(kdb_path,0.58,chunk_size=150)

    def process_content(info):
        pretext = info['pretext']
        if len(pretext) > 300:
            title = sanitize_filename(pretext.split('\n')[0])
            db = llm_kb.get_faiss_db(title, pretext)
            content = llm_kb.search(db, info['content'])
            return content
        return pretext

    #with open(exam_js_path, 'r', encoding='utf-8') as exam_json:
    #    exam_info = json.load(exam_json)

    print("开始进行知识库构建(请关闭VPN)....")

    for info_k, info_v in tqdm(exam_info.items()):
        if info_v['sub_topic']:
            for sub_k, sub_v in info_v['sub_topic'].items():
                exam_info[info_k]['sub_topic'][sub_k]['pretext'] = process_content(sub_v)
        else:
            exam_info[info_k]['pretext'] = process_content(info_v)

    with open(save_js_path, 'w', encoding='utf-8') as json_file:
        json.dump(exam_info, json_file, ensure_ascii=False)
    
    print("知识库已构建完成！")
